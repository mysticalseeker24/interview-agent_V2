"""
Stage 1: Text Extraction Service
Converts any uploaded resume file into plain text.
"""
import logging
import os
from pathlib import Path
from typing import Tuple

import pypdf
from docx import Document

logger = logging.getLogger(__name__)


class TextExtractor:
    """
    Unified text extraction from multiple file formats.
    Strategy: pypdf first, then fallbacks for better coverage.
    """
    
    def __init__(self):
        self.supported_formats = {'.pdf', '.docx', '.doc', '.txt'}
    
    def extract_text(self, file_path: str) -> Tuple[str, str]:
        """
        Extract text from file and return (text, extraction_method).
        
        Args:
            file_path: Path to the resume file
            
        Returns:
            Tuple of (extracted_text, method_used)
        """
        file_path = Path(file_path)
        
        if not file_path.exists():
            raise FileNotFoundError(f"File not found: {file_path}")
        
        file_ext = file_path.suffix.lower()
        
        if file_ext == '.pdf':
            return self._extract_pdf(file_path)
        elif file_ext in ['.docx', '.doc']:
            return self._extract_docx(file_path)
        elif file_ext == '.txt':
            return self._extract_txt(file_path)
        else:
            raise ValueError(f"Unsupported file format: {file_ext}")
    
    def _extract_pdf(self, file_path: Path) -> Tuple[str, str]:
        """Extract text from PDF using pypdf with Tika fallback."""
        try:
            # Try pypdf first (fastest)
            text = self._extract_with_pypdf(file_path)
            
            # Check if extraction was successful
            if len(text.strip()) < 200:
                logger.warning(f"pypdf extracted only {len(text)} chars, trying Tika fallback")
                return self._extract_with_tika(file_path)
            
            logger.info(f"Successfully extracted {len(text)} characters using pypdf")
            return text, "pypdf"
            
        except Exception as e:
            logger.warning(f"pypdf extraction failed: {e}. Trying Tika fallback")
            return self._extract_with_tika(file_path)
    
    def _extract_with_pypdf(self, file_path: Path) -> str:
        """Extract text using pypdf library."""
        text = ""
        with open(file_path, 'rb') as file:
            pdf_reader = pypdf.PdfReader(file)
            for page in pdf_reader.pages:
                text += page.extract_text() + "\n\n"
        return text.strip()
    
    def _extract_with_tika(self, file_path: Path) -> Tuple[str, str]:
        """Extract text using Apache Tika (better layout handling)."""
        try:
            from tika import parser as tika_parser
            result = tika_parser.from_file(str(file_path))
            text = result.get("content", "") if result else ""
            
            if len(text.strip()) < 50:
                raise Exception("Tika extraction too sparse")
            
            logger.info(f"Successfully extracted {len(text)} characters using Tika")
            return text.strip(), "tika"
            
        except ImportError:
            logger.error("Tika not available. Install with: pip install tika")
            raise Exception("Tika fallback not available")
        except Exception as e:
            logger.error(f"Tika extraction failed: {e}")
            raise Exception(f"All PDF extraction methods failed: {e}")
    
    def _extract_docx(self, file_path: Path) -> Tuple[str, str]:
        """Extract text from DOCX file."""
        try:
            doc = Document(file_path)
            text = ""
            
            # Extract paragraphs
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
            
            # Extract tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text += cell.text + " "
                    text += "\n"
            
            logger.info(f"Successfully extracted {len(text)} characters from DOCX")
            return text.strip(), "python-docx"
            
        except Exception as e:
            logger.error(f"DOCX extraction failed: {e}")
            raise Exception(f"Failed to extract DOCX: {e}")
    
    def _extract_txt(self, file_path: Path) -> Tuple[str, str]:
        """Extract text from plain text file."""
        try:
            # Try different encodings
            encodings = ['utf-8', 'latin-1', 'cp1252']
            
            for encoding in encodings:
                try:
                    with open(file_path, 'r', encoding=encoding) as file:
                        text = file.read()
                    logger.info(f"Successfully extracted {len(text)} characters from TXT")
                    return text.strip(), "text-file"
                except UnicodeDecodeError:
                    continue
            
            raise Exception("Could not decode text file with any encoding")
            
        except Exception as e:
            logger.error(f"TXT extraction failed: {e}")
            raise Exception(f"Failed to extract TXT: {e}")
    
    def validate_extraction(self, text: str) -> bool:
        """
        Validate that text extraction was meaningful.
        
        Args:
            text: Extracted text
            
        Returns:
            True if extraction seems valid
        """
        if len(text.strip()) < 100:
            return False
        
        # Check for common resume indicators
        resume_indicators = [
            'experience', 'education', 'skills', 'contact', 'email',
            'phone', 'work', 'project', 'university', 'college'
        ]
        
        text_lower = text.lower()
        found_indicators = sum(1 for indicator in resume_indicators if indicator in text_lower)
        
        return found_indicators >= 3
