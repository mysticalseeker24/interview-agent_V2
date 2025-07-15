"""
Advanced Text Extraction Service
Uses multiple techniques to extract every detail from resumes including hidden links.
"""
import logging
import os
import re
import json
from pathlib import Path
from typing import Tuple, Dict, List, Optional, Any
from dataclasses import dataclass

import pypdf
from docx import Document
import fitz  # PyMuPDF for advanced PDF parsing
import cv2
import numpy as np
from PIL import Image
import pytesseract
from urllib.parse import urlparse, unquote
import io

logger = logging.getLogger(__name__)


@dataclass
class ExtractedElement:
    """Represents an extracted element with metadata."""
    text: str
    element_type: str  # 'text', 'link', 'image', 'table', 'form_field'
    confidence: float
    position: Optional[Tuple[int, int, int, int]] = None  # x, y, width, height
    metadata: Optional[Dict[str, Any]] = None


class AdvancedTextExtractor:
    """
    Advanced text extraction using multiple techniques:
    - OCR for scanned documents
    - Deep PDF parsing for hidden elements
    - Link extraction from annotations and metadata
    - Table and form field extraction
    - Image analysis for embedded text
    """
    
    def __init__(self):
        self.supported_formats = {'.pdf', '.docx', '.doc', '.txt', '.jpg', '.jpeg', '.png'}
        self.extracted_elements = []
        
        # Processing options
        self.processing_options = {
            'enable_ocr': True,
            'enable_hidden_links': True,
            'enable_metadata': True,
            'enable_tables': True,
            'enable_image_ocr': True
        }
    
    def set_processing_options(self, **options):
        """
        Set processing options for text extraction.
        
        Args:
            **options: Processing options to set
                - enable_ocr: Enable OCR for text extraction
                - enable_hidden_links: Enable hidden link detection
                - enable_metadata: Enable metadata extraction
                - enable_tables: Enable table extraction
                - enable_image_ocr: Enable OCR for embedded images
        """
        for key, value in options.items():
            if key in self.processing_options:
                self.processing_options[key] = value
                logger.info(f"Set text extraction option {key}: {value}")
    
    def extract_text(self, file_path: str) -> Tuple[str, str, Dict[str, Any]]:
        """
        Extract text and all elements from file.
        
        Returns:
            Tuple of (extracted_text, method_used, metadata)
        """
        file_path = Path(file_path)
        
        if not file_path.exists():
            raise FileNotFoundError(f"File not found: {file_path}")
        
        file_ext = file_path.suffix.lower()
        
        if file_ext == '.pdf':
            return self._extract_pdf_advanced(file_path)
        elif file_ext in ['.docx', '.doc']:
            return self._extract_docx_advanced(file_path)
        elif file_ext == '.txt':
            return self._extract_txt_advanced(file_path)
        elif file_ext in ['.jpg', '.jpeg', '.png']:
            return self._extract_image_advanced(file_path)
        else:
            raise ValueError(f"Unsupported file format: {file_ext}")
    
    def _extract_pdf_advanced(self, file_path: Path) -> Tuple[str, str, Dict[str, Any]]:
        """Advanced PDF extraction using PyMuPDF with OCR fallback."""
        metadata = {
            'links': [],
            'images': [],
            'tables': [],
            'form_fields': [],
            'annotations': [],
            'hidden_text': [],
            'metadata': {}
        }
        
        try:
            # Use PyMuPDF for comprehensive extraction
            doc = fitz.open(str(file_path))
            full_text = ""
            extraction_method = "pymupdf"
            
            for page_num in range(len(doc)):
                page = doc[page_num]
                
                # Extract visible text
                text = page.get_text()
                full_text += text + "\n\n"
                
                # Extract links and annotations (if enabled)
                if self.processing_options['enable_hidden_links']:
                    links = page.get_links()
                    for link in links:
                        if link.get('uri'):
                            metadata['links'].append({
                                'url': link['uri'],
                                'text': link.get('text', ''),
                                'page': page_num + 1,
                                'rect': link.get('rect', [])
                            })
                    
                    # Extract annotations (comments, highlights, etc.)
                    annotations = page.annots()
                    for annot in annotations:
                        if annot.type[0] == 8:  # Link annotation
                            if annot.uri:
                                metadata['links'].append({
                                    'url': annot.uri,
                                    'text': annot.get_text(),
                                    'page': page_num + 1,
                                    'type': 'annotation'
                                })
                
                # Extract images and check for embedded text (if enabled)
                if self.processing_options['enable_image_ocr']:
                    images = page.get_images()
                    for img_index, img in enumerate(images):
                        try:
                            xref = img[0]
                            pix = fitz.Pixmap(doc, xref)
                            if pix.n - pix.alpha < 4:  # GRAY or RGB
                                img_data = pix.tobytes("png")
                                # Try OCR on image
                                ocr_text = self._ocr_image_from_bytes(img_data)
                                if ocr_text.strip():
                                    metadata['hidden_text'].append({
                                        'text': ocr_text,
                                        'page': page_num + 1,
                                        'source': 'image_ocr'
                                    })
                            pix = None
                        except Exception as e:
                            logger.warning(f"Failed to process image {img_index} on page {page_num}: {e}")
                
                # Extract tables (if enabled)
                if self.processing_options['enable_tables']:
                    tables = page.find_tables()
                    for table in tables:
                        table_data = []
                        for row in table.extract():
                            table_data.append([cell.strip() for cell in row])
                        metadata['tables'].append({
                            'data': table_data,
                            'page': page_num + 1
                        })
                
                # Extract form fields
                widgets = page.widgets()
                for widget in widgets:
                    metadata['form_fields'].append({
                        'name': widget.field_name,
                        'value': widget.field_value,
                        'type': widget.field_type,
                        'page': page_num + 1
                    })
            
            # Extract document metadata (if enabled)
            if self.processing_options['enable_metadata']:
                metadata['metadata'] = {
                    'title': doc.metadata.get('title', ''),
                    'author': doc.metadata.get('author', ''),
                    'subject': doc.metadata.get('subject', ''),
                    'creator': doc.metadata.get('creator', ''),
                    'producer': doc.metadata.get('producer', ''),
                    'creation_date': doc.metadata.get('creationDate', ''),
                    'modification_date': doc.metadata.get('modDate', '')
                }
            
            doc.close()
            
            # If text extraction was poor and OCR is enabled, try OCR
            if len(full_text.strip()) < 200 and self.processing_options['enable_ocr']:
                logger.info("Text extraction poor, trying OCR")
                ocr_text = self._ocr_pdf(file_path)
                if len(ocr_text) > len(full_text):
                    full_text = ocr_text
                    extraction_method = "ocr"
            
            # Add hidden text to main text
            for hidden in metadata['hidden_text']:
                full_text += f"\n{hidden['text']}\n"
            
            # Extract additional links from text
            text_links = self._extract_links_from_text(full_text)
            metadata['links'].extend(text_links)
            
            logger.info(f"Advanced PDF extraction: {len(full_text)} chars, {len(metadata['links'])} links")
            return full_text.strip(), extraction_method, metadata
            
        except Exception as e:
            logger.error(f"Advanced PDF extraction failed: {e}")
            # Fallback to basic extraction
            return self._extract_pdf_basic(file_path)
    
    def _extract_pdf_basic(self, file_path: Path) -> Tuple[str, str, Dict[str, Any]]:
        """Basic PDF extraction as fallback."""
        try:
            with open(file_path, 'rb') as file:
                pdf_reader = pypdf.PdfReader(file)
                text = ""
                for page in pdf_reader.pages:
                    text += page.extract_text() + "\n\n"
                return text.strip(), "pypdf", {'links': [], 'images': [], 'tables': [], 'form_fields': [], 'annotations': [], 'hidden_text': [], 'metadata': {}}
        except Exception as e:
            raise Exception(f"PDF extraction failed: {e}")
    
    def _extract_docx_advanced(self, file_path: Path) -> Tuple[str, str, Dict[str, Any]]:
        """Advanced DOCX extraction with hyperlinks and metadata."""
        metadata = {
            'links': [],
            'images': [],
            'tables': [],
            'form_fields': [],
            'annotations': [],
            'hidden_text': [],
            'metadata': {}
        }
        
        try:
            doc = Document(file_path)
            full_text = ""
            
            # Extract paragraphs with hyperlinks
            for paragraph in doc.paragraphs:
                text = paragraph.text
                full_text += text + "\n"
                
                # Extract hyperlinks from paragraph
                for run in paragraph.runs:
                    if run._element.xpath('.//w:hyperlink'):
                        for hyperlink in run._element.xpath('.//w:hyperlink'):
                            rel_id = hyperlink.get('{http://schemas.openxmlformats.org/officeDocument/2006/relationships}id')
                            if rel_id and rel_id in doc.part.rels:
                                target = doc.part.rels[rel_id].target_ref
                                metadata['links'].append({
                                    'url': target,
                                    'text': run.text,
                                    'type': 'hyperlink'
                                })
            
            # Extract tables
            for table in doc.tables:
                table_data = []
                for row in table.rows:
                    row_data = []
                    for cell in row.cells:
                        row_data.append(cell.text.strip())
                    table_data.append(row_data)
                metadata['tables'].append({'data': table_data})
            
            # Extract document properties
            core_props = doc.core_properties
            metadata['metadata'] = {
                'title': core_props.title or '',
                'author': core_props.author or '',
                'subject': core_props.subject or '',
                'created': str(core_props.created) if core_props.created else '',
                'modified': str(core_props.modified) if core_props.modified else ''
            }
            
            # Extract links from text
            text_links = self._extract_links_from_text(full_text)
            metadata['links'].extend(text_links)
            
            logger.info(f"Advanced DOCX extraction: {len(full_text)} chars, {len(metadata['links'])} links")
            return full_text.strip(), "python-docx", metadata
            
        except Exception as e:
            logger.error(f"Advanced DOCX extraction failed: {e}")
            raise Exception(f"Failed to extract DOCX: {e}")
    
    def _extract_txt_advanced(self, file_path: Path) -> Tuple[str, str, Dict[str, Any]]:
        """Advanced text extraction with link detection."""
        metadata = {
            'links': [],
            'images': [],
            'tables': [],
            'form_fields': [],
            'annotations': [],
            'hidden_text': [],
            'metadata': {}
        }
        
        try:
            # Try different encodings
            encodings = ['utf-8', 'latin-1', 'cp1252', 'utf-16']
            
            for encoding in encodings:
                try:
                    with open(file_path, 'r', encoding=encoding) as file:
                        text = file.read()
                    
                    # Extract links from text
                    text_links = self._extract_links_from_text(text)
                    metadata['links'].extend(text_links)
                    
                    logger.info(f"Advanced TXT extraction: {len(text)} chars, {len(metadata['links'])} links")
                    return text.strip(), "text-file", metadata
                except UnicodeDecodeError:
                    continue
            
            raise Exception("Could not decode text file with any encoding")
            
        except Exception as e:
            logger.error(f"Advanced TXT extraction failed: {e}")
            raise Exception(f"Failed to extract TXT: {e}")
    
    def _extract_image_advanced(self, file_path: Path) -> Tuple[str, str, Dict[str, Any]]:
        """Extract text from images using OCR."""
        metadata = {
            'links': [],
            'images': [],
            'tables': [],
            'form_fields': [],
            'annotations': [],
            'hidden_text': [],
            'metadata': {}
        }
        
        try:
            # Use OCR to extract text from image
            text = self._ocr_image(str(file_path))
            
            # Extract links from OCR text
            text_links = self._extract_links_from_text(text)
            metadata['links'].extend(text_links)
            
            logger.info(f"Advanced image extraction: {len(text)} chars, {len(metadata['links'])} links")
            return text.strip(), "ocr", metadata
            
        except Exception as e:
            logger.error(f"Advanced image extraction failed: {e}")
            raise Exception(f"Failed to extract image: {e}")
    
    def _ocr_pdf(self, file_path: Path) -> str:
        """OCR a PDF file."""
        try:
            doc = fitz.open(str(file_path))
            text = ""
            
            for page_num in range(len(doc)):
                page = doc[page_num]
                
                # Convert page to image
                pix = page.get_pixmap()
                img_data = pix.tobytes("png")
                
                # OCR the image
                page_text = self._ocr_image_from_bytes(img_data)
                text += page_text + "\n\n"
                
                pix = None
            
            doc.close()
            return text
            
        except Exception as e:
            logger.error(f"PDF OCR failed: {e}")
            return ""
    
    def _ocr_image(self, image_path: str) -> str:
        """OCR an image file."""
        try:
            # Preprocess image for better OCR
            image = cv2.imread(image_path)
            if image is None:
                raise Exception("Could not load image")
            
            # Convert to grayscale
            gray = cv2.cvtColor(image, cv2.COLOR_BGR2GRAY)
            
            # Apply preprocessing for better OCR
            # Noise reduction
            denoised = cv2.fastNlMeansDenoising(gray)
            
            # Thresholding
            _, thresh = cv2.threshold(denoised, 0, 255, cv2.THRESH_BINARY + cv2.THRESH_OTSU)
            
            # OCR with multiple configurations
            configs = [
                '--oem 3 --psm 6',  # Default
                '--oem 3 --psm 3',  # Fully automatic page segmentation
                '--oem 1 --psm 6',  # Legacy engine
            ]
            
            best_text = ""
            for config in configs:
                try:
                    text = pytesseract.image_to_string(thresh, config=config)
                    if len(text.strip()) > len(best_text.strip()):
                        best_text = text
                except Exception:
                    continue
            
            return best_text
            
        except Exception as e:
            logger.error(f"Image OCR failed: {e}")
            return ""
    
    def _ocr_image_from_bytes(self, image_bytes: bytes) -> str:
        """OCR an image from bytes."""
        try:
            # Convert bytes to PIL Image
            image = Image.open(io.BytesIO(image_bytes))
            
            # Convert to OpenCV format
            opencv_image = cv2.cvtColor(np.array(image), cv2.COLOR_RGB2BGR)
            
            # Preprocess
            gray = cv2.cvtColor(opencv_image, cv2.COLOR_BGR2GRAY)
            denoised = cv2.fastNlMeansDenoising(gray)
            _, thresh = cv2.threshold(denoised, 0, 255, cv2.THRESH_BINARY + cv2.THRESH_OTSU)
            
            # OCR
            text = pytesseract.image_to_string(thresh, config='--oem 3 --psm 6')
            return text
            
        except Exception as e:
            logger.error(f"Image bytes OCR failed: {e}")
            return ""
    
    def _extract_links_from_text(self, text: str) -> List[Dict[str, str]]:
        """Extract all types of links from text."""
        links = []
        
        # URL patterns
        url_patterns = [
            r'https?://[^\s<>"{}|\\^`\[\]]+',
            r'www\.[^\s<>"{}|\\^`\[\]]+',
            r'ftp://[^\s<>"{}|\\^`\[\]]+',
            r'mailto:[^\s<>"{}|\\^`\[\]]+',
        ]
        
        # Social media patterns
        social_patterns = [
            r'linkedin\.com/in/([a-zA-Z0-9-]+)',
            r'github\.com/([a-zA-Z0-9-]+)',
            r'twitter\.com/([a-zA-Z0-9_]+)',
            r'facebook\.com/([a-zA-Z0-9.]+)',
            r'instagram\.com/([a-zA-Z0-9._]+)',
        ]
        
        # Extract URLs
        for pattern in url_patterns:
            matches = re.finditer(pattern, text, re.IGNORECASE)
            for match in matches:
                url = match.group(0)
                # Get surrounding context
                start = max(0, match.start() - 50)
                end = min(len(text), match.end() + 50)
                context = text[start:end].strip()
                
                links.append({
                    'url': url,
                    'text': context,
                    'type': 'url',
                    'position': (match.start(), match.end())
                })
        
        # Extract social media profiles
        for pattern in social_patterns:
            matches = re.finditer(pattern, text, re.IGNORECASE)
            for match in matches:
                platform = pattern.split('.')[0]
                username = match.group(1)
                url = f"https://{match.group(0)}"
                
                links.append({
                    'url': url,
                    'text': f"{platform} profile: {username}",
                    'type': 'social_media',
                    'platform': platform,
                    'username': username,
                    'position': (match.start(), match.end())
                })
        
        # Extract email addresses
        email_pattern = r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b'
        email_matches = re.finditer(email_pattern, text)
        for match in email_matches:
            email = match.group(0)
            links.append({
                'url': f"mailto:{email}",
                'text': email,
                'type': 'email',
                'position': (match.start(), match.end())
            })
        
        return links
    
    def validate_extraction(self, text: str, metadata: Dict[str, Any]) -> bool:
        """
        Validate that extraction was meaningful.
        
        Args:
            text: Extracted text
            metadata: Extraction metadata
            
        Returns:
            True if extraction seems valid
        """
        if len(text.strip()) < 50:
            return False
        
        # Check for common resume indicators
        resume_indicators = [
            'experience', 'education', 'skills', 'contact', 'email',
            'phone', 'work', 'project', 'university', 'college',
            'resume', 'cv', 'curriculum vitae'
        ]
        
        text_lower = text.lower()
        found_indicators = sum(1 for indicator in resume_indicators if indicator in text_lower)
        
        # Also check if we found any links or useful metadata
        has_links = len(metadata.get('links', [])) > 0
        has_metadata = any(metadata.get('metadata', {}).values())
        
        return found_indicators >= 2 or has_links or has_metadata
    
    def get_extraction_summary(self, text: str, metadata: Dict[str, Any]) -> Dict[str, Any]:
        """Get a summary of the extraction results."""
        return {
            'text_length': len(text),
            'links_found': len(metadata.get('links', [])),
            'tables_found': len(metadata.get('tables', [])),
            'images_found': len(metadata.get('images', [])),
            'form_fields_found': len(metadata.get('form_fields', [])),
            'annotations_found': len(metadata.get('annotations', [])),
            'hidden_text_found': len(metadata.get('hidden_text', [])),
            'metadata_fields': len([v for v in metadata.get('metadata', {}).values() if v]),
            'link_types': list(set(link.get('type', 'unknown') for link in metadata.get('links', []))),
            'social_media_profiles': [
                link for link in metadata.get('links', []) 
                if link.get('type') == 'social_media'
            ]
        } 